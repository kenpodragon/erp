"""
Phase 1: Text Extraction & Structural Splitting.

Steps per chapter:
  1. Read the .docx file and segment text into chapters (by heading style).
  2. Detect hard breaks within each chapter to get candidate scene segments.
  3. Send each chapter's segments to AI to produce final scene + beat boundaries.
  4. Persist everything to DB: chapters, scenes, story_beats.
  5. Mark chapter processing_status = 'text_extracted'.
"""
from __future__ import annotations
import json
import logging
from pathlib import Path
from typing import Optional

import docx
from docx import Document

from config import (
    BOOK_REGISTRY, BOOKS_DIR, CHAPTER_HEADING_STYLES, HARD_BREAK_PATTERNS,
    BOOK_BY_NUMBER,
)
from core.ai_provider import AIProvider
from core.db import (
    get_or_create_book, upsert_chapter, upsert_scene, upsert_story_beat,
    get_chapters_for_book, update_chapter_status, transaction,
)
from core.display import ProgressTracker
from core.models import Phase1Response, SceneStructure, BeatStructure

logger = logging.getLogger("book_parser.phase1")

# ── System prompt ─────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a literary structure analyst. Your task is to split a book chapter
into scenes and story beats based on the provided paragraph text.

DEFINITIONS:
- Scene: A coherent narrative unit with a consistent location and set of characters.
  Scenes change when: location shifts, significant time passes, or a hard break marker (***) appears.
- Story Beat: A single logical moment within a scene — one action, exchange, or shift in tension.
  Typically 1–5 paragraphs. Keep beats cohesive; avoid splitting mid-action.

GUIDELINES:
- A chapter must have at least 1 scene.
- Each scene must have at least 1 beat.
- Prefer 2–5 beats per scene when the text supports it.
- If a hard break marker was detected (triggered_by_hard_break=true), honour that as an absolute scene boundary.
- For pacing use one of: slow-burn | action | dialogue-heavy | contemplative | transitional
- For timeline_context use: present | flashback | dream | vision | future
- Intensity is 1 (calm) to 5 (peak tension).
- Paragraph indices are 0-based.
"""


def _build_user_prompt(chapter_title: str, chapter_number: int,
                       paragraphs: list[str], hard_break_indices: list[int]) -> str:
    """Build the user-facing prompt containing the chapter text."""
    para_text = "\n".join(
        f"[{i}] {p}" for i, p in enumerate(paragraphs)
    )
    break_info = (
        f"Hard break markers detected before paragraphs: {hard_break_indices}"
        if hard_break_indices
        else "No hard break markers detected."
    )
    return (
        f"Chapter {chapter_number}: {chapter_title or '(untitled)'}\n\n"
        f"{break_info}\n\n"
        f"PARAGRAPHS:\n{para_text}\n\n"
        "Produce the scene and beat structure for this chapter."
    )


# ── .docx parsing ─────────────────────────────────────────────────────────

def _is_chapter_heading(para) -> bool:
    return para.style.name in CHAPTER_HEADING_STYLES


def _is_hard_break(para) -> bool:
    text = para.text.strip()
    return text in HARD_BREAK_PATTERNS


def _para_text(para) -> str:
    return para.text.strip()


def _extract_chapters_from_docx(docx_path: Path) -> list[dict]:
    """
    Parse the .docx and return a list of chapter dicts:
    {
        "chapter_number": int,
        "title": str | None,
        "paragraphs": [str, ...],           # raw text, hard-break lines removed
        "hard_break_indices": [int, ...],   # para indices where a hard break preceded
    }
    """
    doc = Document(str(docx_path))
    chapters: list[dict] = []
    current_title: Optional[str] = None
    current_paras: list[str] = []
    hard_breaks: list[int] = []

    def _flush() -> None:
        nonlocal current_title, current_paras, hard_breaks
        if current_paras:
            chapters.append({
                "chapter_number": len(chapters) + 1,
                "title": current_title,
                "paragraphs": current_paras,
                "hard_break_indices": hard_breaks,
            })
        current_title = None
        current_paras = []
        hard_breaks = []

    for para in doc.paragraphs:
        if _is_chapter_heading(para):
            _flush()
            current_title = para.text.strip() or None
        elif _is_hard_break(para):
            # Record where the next paragraph index will be
            hard_breaks.append(len(current_paras))
        else:
            text = _para_text(para)
            if text:
                current_paras.append(text)

    _flush()

    # If no heading styles found, treat the whole file as one chapter
    if not chapters:
        all_paras: list[str] = []
        all_breaks: list[int] = []
        for para in doc.paragraphs:
            if _is_hard_break(para):
                all_breaks.append(len(all_paras))
            else:
                text = _para_text(para)
                if text:
                    all_paras.append(text)
        if all_paras:
            chapters.append({
                "chapter_number": 1,
                "title": None,
                "paragraphs": all_paras,
                "hard_break_indices": all_breaks,
            })

    return chapters


# ── Scene / beat building from AI response ────────────────────────────────

def _extract_raw_text(paragraphs: list[str], start: int, end: int) -> str:
    """Slice paragraphs [start..end] (inclusive) into a single text block."""
    start = max(0, start)
    end = min(len(paragraphs) - 1, end)
    return "\n\n".join(paragraphs[start:end + 1])


def _persist_chapter(conn, book_id: int, ch_data: dict, ai_response: Phase1Response,
                     paragraphs: list[str]) -> int:
    """Write chapter + all scenes + all beats to DB. Returns chapter_id."""
    raw_text = "\n\n".join(paragraphs)
    chapter_id = upsert_chapter(
        conn,
        book_id=book_id,
        chapter_number=ch_data["chapter_number"],
        title=ch_data["title"],
        raw_text=raw_text,
        sort_order=ch_data["chapter_number"],
    )

    for scene_idx, scene in enumerate(ai_response.scenes):
        scene_raw = _extract_raw_text(paragraphs, scene.start_para_idx, scene.end_para_idx)
        scene_id = upsert_scene(
            conn,
            chapter_id=chapter_id,
            scene_number=scene.scene_number,
            title=scene.title,
            summary=scene.summary,
            raw_text=scene_raw,
            sort_order=scene_idx,
            has_hard_break=scene.triggered_by_hard_break,
        )

        for beat_idx, beat in enumerate(scene.beats):
            beat_raw = _extract_raw_text(paragraphs, beat.start_para_idx, beat.end_para_idx)
            upsert_story_beat(
                conn,
                scene_id=scene_id,
                beat_number=beat.beat_number,
                summary=beat.summary,
                raw_text=beat_raw,
                sort_order=beat_idx,
                intensity=beat.intensity,
                pacing=beat.pacing,
                timeline_context=beat.timeline_context,
            )

    return chapter_id


# ── Main phase entry point ────────────────────────────────────────────────

def run_phase1(conn, book_number: int, ai: AIProvider,
               tracker: ProgressTracker,
               resume_after_chapter: Optional[int] = None) -> None:
    """
    Run Phase 1 for a single book.

    Args:
        conn: psycopg2 connection (caller manages commit/rollback).
        book_number: The book to process (1, 2, or 3).
        ai: Configured AIProvider instance.
        tracker: ProgressTracker for live display updates.
        resume_after_chapter: If set, skip chapters ≤ this number.
    """
    book_meta = BOOK_BY_NUMBER.get(book_number)
    if not book_meta:
        raise ValueError(f"Unknown book number: {book_number}")

    docx_path = BOOKS_DIR / book_meta["source_file"]
    if not docx_path.exists():
        raise FileNotFoundError(f"Book file not found: {docx_path}")

    logger.info("Phase 1 — Book %d: %s (%s)", book_number, book_meta["title"], docx_path.name)

    # Ensure book row exists
    book_id = get_or_create_book(conn, book_number, book_meta["title"], book_meta["source_file"])
    with transaction(conn):
        pass  # commit book creation

    # Parse docx
    logger.info("Parsing %s ...", docx_path.name)
    chapters = _extract_chapters_from_docx(docx_path)
    logger.info("Found %d chapters in %s", len(chapters), docx_path.name)

    tracker.set_book(book_number, book_meta["title"], phase=1)
    tracker.set_chapters(len(chapters))

    for ch_data in chapters:
        ch_num = ch_data["chapter_number"]

        # Resume: skip already-processed chapters
        if resume_after_chapter is not None and ch_num <= resume_after_chapter:
            logger.debug("Skipping chapter %d (already processed)", ch_num)
            tracker.advance_chapter()
            continue

        ch_title = ch_data["title"] or f"Chapter {ch_num}"
        tracker.advance_chapter(ch_title)
        logger.info("  Processing chapter %d: %s", ch_num, ch_title)

        paragraphs = ch_data["paragraphs"]
        if not paragraphs:
            logger.warning("  Chapter %d has no paragraphs — skipping", ch_num)
            continue

        user_prompt = _build_user_prompt(
            ch_title, ch_num, paragraphs, ch_data["hard_break_indices"]
        )

        tracker.set_items("Calling AI for scene/beat structure...", total=1)

        try:
            json_text, provider, model_id = ai.call(
                system_prompt=SYSTEM_PROMPT,
                user_prompt=user_prompt,
                response_model=Phase1Response,
            )
        except Exception as e:
            logger.error("  AI call failed for chapter %d: %s", ch_num, e)
            raise

        tracker.advance_item()

        phase1_result: Phase1Response = ai.parse_response(json_text, Phase1Response)
        logger.info(
            "  Chapter %d → %d scenes, %d total beats (via %s)",
            ch_num,
            len(phase1_result.scenes),
            sum(len(s.beats) for s in phase1_result.scenes),
            provider,
        )

        # Persist within its own transaction so a single bad chapter doesn't abort the book
        with transaction(conn):
            chapter_id = _persist_chapter(conn, book_id, ch_data, phase1_result, paragraphs)
            update_chapter_status(conn, chapter_id, "text_extracted")

        tracker.log(
            f"Chapter {ch_num} ({ch_title}): "
            f"{len(phase1_result.scenes)} scenes, "
            f"{sum(len(s.beats) for s in phase1_result.scenes)} beats  [{provider}]",
            style="green",
        )

    logger.info("Phase 1 complete for Book %d", book_number)
